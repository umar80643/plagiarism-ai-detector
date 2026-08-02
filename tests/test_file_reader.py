import io

import docx
import pytest
from pypdf import PdfWriter

from utils.file_reader import UnsupportedFileTypeError, extract_text


def test_extract_text_from_txt_path(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("hello from a text file", encoding="utf-8")
    assert extract_text(path) == "hello from a text file"


def test_extract_text_from_txt_file_like_object_with_name():
    class FakeUpload:
        name = "sample.txt"

        def read(self):
            return b"hello from an upload"

    assert extract_text(FakeUpload()) == "hello from an upload"


def test_extract_text_from_docx(tmp_path):
    path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("first paragraph")
    document.add_paragraph("second paragraph")
    document.save(path)
    assert extract_text(path) == "first paragraph\nsecond paragraph"


def test_extract_text_from_pdf(tmp_path):
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as handle:
        writer.write(handle)
    # A blank page has no extractable text, but this should not raise.
    assert extract_text(path) == ""


def test_extract_text_rejects_unsupported_extension(tmp_path):
    path = tmp_path / "sample.exe"
    path.write_bytes(b"not text")
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(path)
